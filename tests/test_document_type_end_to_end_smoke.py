from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from app.extractors.equipment_extractor import extract_equipment_records
from app.extractors.points_list_extractor import extract_point_records
from app.extractors.template_fallback_merge import merge_point_records_with_template_defaults
from app.normalizers.markdown_normalizer import normalize_text
from app.parsers.docx_parser import parse_docx_to_markdown
from app.parsers.excel_parser import parse_excel_to_markdown
from app.parsers.pdf_inspector_parser import inspect_pdf


SAMPLE_TEXT = """## Equipment Schedule

| Tag | Type | Equipment | Manufacturer | Model | Voltage | CFM | Protocol |
|---|---|---|---|---|---|---:|---|
| AHU-1 | AHU | Air Handling Unit 1 | Daikin | Vision | 480V/3PH | 12000 | BACnet/IP |
| RTU-2 | RTU | Roof Top Unit 2 | Trane | Voyager | 480V/3PH | 4200 | BACnet/IP |

## Controls Points

| Equipment Tag | Point Name | Abbr | IO Type | Unit | Description |
|---|---|---|---|---|---|
| AHU-1 | Supply Air Temperature | SAT | AI | degF | AHU-1 discharge air sensor |
| RTU-2 | Space Temperature | ZN-T | AI | degF | RTU-2 space temperature sensor |
"""


class DocumentTypeEndToEndSmokeTests(unittest.TestCase):
    def _write_simple_pdf(self, path: Path, text: str) -> None:
        lines = [line for line in text.splitlines() if line.strip()]

        def esc(value: str) -> str:
            return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        text_lines = ["BT", "/F1 11 Tf", "72 800 Td", "14 TL"]
        for index, line in enumerate(lines):
            operator = "Tj" if index == 0 else "'"
            text_lines.append(f"({esc(line)}) {operator}")
        text_lines.append("ET")
        stream = "\n".join(text_lines).encode("utf-8")

        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"<< /Length %d >>\nstream\n%b\nendstream" % (len(stream), stream),
        ]

        content = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: list[int] = []
        for index, obj in enumerate(objects, start=1):
            offsets.append(len(content))
            content.extend(f"{index} 0 obj\n".encode("ascii"))
            content.extend(obj)
            content.extend(b"\nendobj\n")

        xref_offset = len(content)
        content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        content.extend(b"0000000000 65535 f \n")
        for offset in offsets:
            content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        content.extend(
            (
                f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF\n"
            ).encode("ascii")
        )
        path.write_bytes(content)

    def _run_pipeline(self, raw_markdown: str) -> tuple[list, list, list]:
        normalized = normalize_text(raw_markdown)
        equipment = extract_equipment_records(normalized)
        points = extract_point_records(normalized, equipment_records=equipment)
        merged_points = merge_point_records_with_template_defaults(points, equipment)
        return equipment, points, merged_points

    def _assert_pipeline_outputs(self, equipment: list, points: list, merged_points: list) -> None:
        self.assertEqual(len(equipment), 2)
        self.assertEqual({record.equipment_tag for record in equipment}, {"AHU-1", "RTU-2"})
        self.assertEqual(len(points), 2)
        self.assertTrue(all(point.equipment_tag in {"AHU-1", "RTU-2"} for point in points))
        self.assertGreater(len(merged_points), len(points))
        self.assertTrue(any(point.source_type == "template_default" for point in merged_points))

    def test_end_to_end_pasted_text(self) -> None:
        equipment, points, merged_points = self._run_pipeline(SAMPLE_TEXT)
        self._assert_pipeline_outputs(equipment, points, merged_points)

    def test_end_to_end_docx(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory(prefix="specflow_docx_smoke_") as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.add_heading("Equipment Schedule", level=2)
            equipment_table = doc.add_table(rows=1, cols=8)
            headers = ["Tag", "Type", "Equipment", "Manufacturer", "Model", "Voltage", "CFM", "Protocol"]
            for index, value in enumerate(headers):
                equipment_table.rows[0].cells[index].text = value
            equipment_rows = [
                ["AHU-1", "AHU", "Air Handling Unit 1", "Daikin", "Vision", "480V/3PH", "12000", "BACnet/IP"],
                ["RTU-2", "RTU", "Roof Top Unit 2", "Trane", "Voyager", "480V/3PH", "4200", "BACnet/IP"],
            ]
            for row in equipment_rows:
                cells = equipment_table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value

            doc.add_heading("Controls Points", level=2)
            point_table = doc.add_table(rows=1, cols=6)
            point_headers = ["Equipment Tag", "Point Name", "Abbr", "IO Type", "Unit", "Description"]
            for index, value in enumerate(point_headers):
                point_table.rows[0].cells[index].text = value
            point_rows = [
                ["AHU-1", "Supply Air Temperature", "SAT", "AI", "degF", "AHU-1 discharge air sensor"],
                ["RTU-2", "Space Temperature", "ZN-T", "AI", "degF", "RTU-2 space temperature sensor"],
            ]
            for row in point_rows:
                cells = point_table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value

            doc.save(path)

            result = parse_docx_to_markdown(str(path))
            self.assertEqual(result["status"], "success")

            equipment, points, merged_points = self._run_pipeline(result["raw_markdown"])
            self._assert_pipeline_outputs(equipment, points, merged_points)

    def test_end_to_end_excel(self) -> None:
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory(prefix="specflow_xlsx_smoke_") as tmpdir:
            path = Path(tmpdir) / "sample.xlsx"
            workbook = Workbook()
            equipment_sheet = workbook.active
            equipment_sheet.title = "Equipment Schedule"
            equipment_sheet.append(["Tag", "Type", "Equipment", "Manufacturer", "Model", "Voltage", "CFM", "Protocol"])
            equipment_sheet.append(["AHU-1", "AHU", "Air Handling Unit 1", "Daikin", "Vision", "480V/3PH", 12000, "BACnet/IP"])
            equipment_sheet.append(["RTU-2", "RTU", "Roof Top Unit 2", "Trane", "Voyager", "480V/3PH", 4200, "BACnet/IP"])

            point_sheet = workbook.create_sheet("Controls Points")
            point_sheet.append(["Equipment Tag", "Point Name", "Abbr", "IO Type", "Unit", "Description"])
            point_sheet.append(["AHU-1", "Supply Air Temperature", "SAT", "AI", "degF", "AHU-1 discharge air sensor"])
            point_sheet.append(["RTU-2", "Space Temperature", "ZN-T", "AI", "degF", "RTU-2 space temperature sensor"])
            workbook.save(path)

            result = parse_excel_to_markdown(str(path))
            self.assertEqual(result["status"], "success")

            equipment, points, merged_points = self._run_pipeline(result["raw_markdown"])
            self._assert_pipeline_outputs(equipment, points, merged_points)

    def test_end_to_end_pdf_text_based(self) -> None:
        with tempfile.TemporaryDirectory(prefix="specflow_pdf_smoke_") as tmpdir:
            path = Path(tmpdir) / "sample.pdf"
            self._write_simple_pdf(path, SAMPLE_TEXT.replace("degF", "F"))

            result = inspect_pdf(str(path))
            self.assertEqual(result["status"], "success")
            self.assertIn(result["pdf_classification"], {"text_pdf", "mixed_pdf"})

            equipment, points, merged_points = self._run_pipeline(result["raw_markdown"])
            self._assert_pipeline_outputs(equipment, points, merged_points)


if __name__ == "__main__":
    unittest.main()
